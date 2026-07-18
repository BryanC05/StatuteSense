#!/usr/bin/env python3
"""Export analysis results to DOCX format.
Usage: uv run scripts/export-docx.py --title "Analysis" --content "Markdown text..." --output result.docx
"""
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_to_docx(title: str, content: str, output: str):
    doc = Document()
    
    # Title
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Content (simple markdown-like parsing)
    for line in content.split('\n'):
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.strip():
            p = doc.add_paragraph(line)
            p.paragraph_format.line_spacing = Pt(14)
        else:
            doc.add_paragraph('')
    
    # Save
    doc.save(output)
    print(f"Exported to {output}")

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Export to DOCX')
    parser.add_argument('--title', required=True, help='Document title')
    parser.add_argument('--content', required=True, help='Content text')
    parser.add_argument('--output', default='export.docx', help='Output filename')
    args = parser.parse_args()
    export_to_docx(args.title, args.content, args.output)
